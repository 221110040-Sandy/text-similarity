# app_frontend.py - Streamlit Frontend for Text Similarity
import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import requests
import time
import re
from typing import Dict, Optional

# PDF processing imports
try:
    import PyPDF2
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# DOCX processing imports
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# Set page config
st.set_page_config(
    page_title="📝 Text Similarity Analyzer",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
            
    .main-header p {
        margin-bottom: 0;
    }
    
    .metric-card {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 1rem;
        border-radius: 10px;
        border-left: 5px solid #667eea;
        margin: 0.5rem 0;
    }
    
    .similarity-high {
        background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
        border-left-color: #28a745;
    }
    
    .similarity-medium {
        background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
        border-left-color: #ffc107;
    }
    
    .similarity-low {
        background: linear-gradient(135deg, #f8d7da 0%, #f1c0c7 100%);
        border-left-color: #dc3545;
    }
    
    .speed-indicator {
        padding: 0.25rem 0.5rem;
        border-radius: 12px;
        font-size: 0.8rem;
        font-weight: bold;
        margin-left: 0.5rem;
    }
    
    .speed-fast { background: #d4edda; color: #155724; }
    .speed-medium { background: #fff3cd; color: #856404; }
    .speed-slow { background: #f8d7da; color: #721c24; }
    
    .api-status {
        padding: 0.5rem 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        border-left: 4px solid;
    }
    
    .api-connected { background: #d4edda; border-left-color: #28a745; color: #155724; }
    .api-disconnected { background: #f8d7da; border-left-color: #dc3545; color: #721c24; }
            
    [data-testid="stHeading"] a {
        display: none !important;
    }
            
    h1 a, h2 a, h3 a, h4 a, h5 a, h6 a {
        display: none !important;
    }
</style>
""", unsafe_allow_html=True)

# API Configuration
# API_BASE_URL = "http://localhost:8000"
API_BASE_URL = "https://text-similarity-production.up.railway.app"
# Utility functions
def count_words(text):
    return len(str(text).split())

def estimate_pages(text, words_per_page=250):
    return round(count_words(text) / words_per_page, 1)

def validate_document_length(text, max_pages=5):
    estimated_pages = estimate_pages(text)
    word_count = count_words(text)
    if estimated_pages > max_pages:
        return False, f"❌ Dokumen terlalu panjang! ({estimated_pages} halaman, {word_count} kata). Maksimal {max_pages} halaman (~{max_pages * 250} kata)."
    else:
        return True, f"✅ Dokumen valid ({estimated_pages} halaman, {word_count} kata)"

def get_text_statistics(text):
    word_count = count_words(text)
    char_count = len(text)
    sentence_count = len(re.findall(r'[.!?]+', text))
    paragraph_count = len([p for p in text.split('\n\n') if p.strip()])
    
    # Simplified readability
    avg_words_per_sentence = word_count / max(sentence_count, 1)
    complexity_score = min(100, max(0, 100 - (avg_words_per_sentence * 2)))
    
    return {
        'words': word_count,
        'characters': char_count,
        'sentences': sentence_count,
        'paragraphs': paragraph_count,
        'pages': estimate_pages(text),
        'complexity_score': round(complexity_score, 1),
        'avg_words_per_sentence': round(avg_words_per_sentence, 1)
    }

# PDF Processing Functions
def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file using multiple methods."""
    if not PDF_SUPPORT:
        return None, "📦 PDF support tidak tersedia. Install PyPDF2 dan pdfplumber terlebih dahulu."
    
    try:
        # Method 1: Try pdfplumber first (better for complex layouts)
        text = ""
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            return text.strip(), None
            
        # Method 2: Fallback to PyPDF2
        pdf_file.seek(0)  # Reset file pointer
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
            
        if text.strip():
            return text.strip(), None
        else:
            return None, "❌ Tidak bisa extract teks dari PDF. File mungkin berupa gambar atau rusak."
            
    except Exception as e:
        return None, f"❌ Error memproses PDF: {str(e)}"

def create_pdf_preview(text, max_chars=500):
    """Create a preview of PDF content."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."

def validate_pdf_file(uploaded_file):
    """Validate if uploaded file is a valid PDF."""
    if uploaded_file.type != "application/pdf":
        return False, "❌ File harus berformat PDF"
    
    # Check file size (max 10MB)
    if uploaded_file.size > 10 * 1024 * 1024:
        return False, "❌ File PDF terlalu besar! Maksimal 10MB"
    
    return True, "✅ File PDF valid"

# DOCX Processing Functions
def extract_text_from_docx(docx_file):
    """Extract text from DOCX file."""
    if not DOCX_SUPPORT:
        return None, "📦 DOCX support tidak tersedia. Install python-docx terlebih dahulu."
    
    try:
        doc = Document(docx_file)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        if text.strip():
            return text.strip(), None
        else:
            return None, "❌ Tidak ada teks yang bisa diextract dari DOCX. File mungkin kosong."
            
    except Exception as e:
        return None, f"❌ Error memproses DOCX: {str(e)}"

def create_docx_preview(text, max_chars=500):
    """Create a preview of DOCX content."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."

def validate_docx_file(uploaded_file):
    """Validate if uploaded file is a valid DOCX."""
    # Check MIME type
    valid_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ]
    
    if uploaded_file.type not in valid_types:
        # Also check by file extension
        if not uploaded_file.name.lower().endswith(('.docx', '.doc')):
            return False, "❌ File harus berformat DOCX atau DOC"
    
    # Check file size (max 10MB)
    if uploaded_file.size > 10 * 1024 * 1024:
        return False, "❌ File DOCX terlalu besar! Maksimal 10MB"
    
    return True, "✅ File DOCX valid"

# API Communication Functions
@st.cache_data(ttl=60)  # Cache for 1 minute
def check_api_health():
    """Check if API is running and model is loaded"""
    try:
        response = requests.get(f"{API_BASE_URL}/health", timeout=5)
        if response.status_code == 200:
            return response.json()
        else:
            return {"status": "error", "model_loaded": False}
    except requests.exceptions.RequestException:
        return {"status": "disconnected", "model_loaded": False}

def predict_similarity_api(text1: str, text2: str):
    """Call API for similarity prediction menggunakan Neural model"""
    try:
        payload = {
            "text1": text1,
            "text2": text2
        }
        
        response = requests.post(
            f"{API_BASE_URL}/predict",
            json=payload,
            timeout=30  # 30 second timeout
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            error_detail = response.json().get("detail", "Unknown error")
            return {"error": f"API Error: {error_detail}"}
            
    except requests.exceptions.Timeout:
        return {"error": "Request timeout - API took too long to respond"}
    except requests.exceptions.RequestException as e:
        return {"error": f"Connection error: {str(e)}"}

def predict_document_api(doc1: str, doc2: str, per_side_len: int = 28, stride: int = 28, topk_evidence: int = 5, use_symmetric: bool = True):
    """Call API for long-document similarity (sliding window + BERTScore-like)."""
    try:
        payload = {
            "doc1": doc1,
            "doc2": doc2,
            "per_side_len": per_side_len,
            "stride": stride,
            "topk_evidence": topk_evidence,
            "use_symmetric": use_symmetric
        }
        resp = requests.post(
            f"{API_BASE_URL}/predict-document",
            json=payload,
            timeout=120   # dokumen bisa lebih lama, kasih timeout lebih besar
        )
        if resp.status_code == 200:
            return resp.json()
        else:
            # backend FastAPI biasanya kirim {"detail": "..."}
            det = resp.json().get("detail", f"HTTP {resp.status_code}")
            return {"error": f"API Error: {det}"}
    except requests.exceptions.Timeout:
        return {"error": "Request timeout - document API took too long to respond"}
    except requests.exceptions.RequestException as e:
        return {"error": f"Connection error: {str(e)}"}


# Visualization Functions
def create_similarity_gauge(similarity, method="", processing_time=0):
    # Tentukan warna bar berdasarkan similarity
    if similarity >= 0.8:
        bar_color = "#f5576c"  # Merah untuk plagiarisme tinggi
    elif similarity >= 0.6:
        bar_color = "#fdcb6e"  # Kuning untuk sedang
    else:
        bar_color = "#00b894"  # Hijau untuk rendah/aman
    
    fig = go.Figure(go.Indicator(
        mode = "gauge+number",
        value = similarity * 100,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {
            'text': f"<b style='font-size: 18px; color: #2d3748;'>Similarity Score (%)</b><br><span style='color: #667eea; font-size: 14px; font-weight: 600;'>{method}</span><br><span style='color: #764ba2; font-size: 13px; font-weight: 500;'>⚡ {processing_time:.3f}s</span>",
            'font': {'size': 16, 'color': '#2d3748', 'family': 'Inter, -apple-system, BlinkMacSystemFont, sans-serif'}
        },
        number = {
            'font': {'size': 56, 'color': bar_color, 'family': 'Inter, -apple-system, BlinkMacSystemFont, sans-serif'},
            'suffix': '%'
        },
        gauge = {
            'axis': {
                'range': [None, 100], 
                'tickwidth': 2,
                'tickcolor': '#cbd5e0',
                'tickfont': {'size': 13, 'color': '#4a5568', 'family': 'Inter, sans-serif'}
            },
            'bar': {'color': bar_color, 'thickness': 0.35},
            'bgcolor': 'white',
            'borderwidth': 2,
            'bordercolor': '#e2e8f0',
            'steps': [
                {'range': [0, 60], 'color': 'rgba(0, 184, 148, 0.15)'},   # Hijau muda
                {'range': [60, 80], 'color': 'rgba(253, 203, 110, 0.15)'}, # Kuning muda
                {'range': [80, 100], 'color': 'rgba(245, 87, 108, 0.15)'}  # Merah muda
            ],
            'threshold': {
                'line': {'color': bar_color, 'width': 5},
                'thickness': 0.8,
                'value': similarity * 100
            }
        }
    ))
    
    fig.update_layout(
        height=350, 
        font={'color': "#2d3748", 'family': "Inter, sans-serif"},
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        margin=dict(l=20, r=20, t=80, b=20)
    )
    return fig

def create_performance_chart(results_history):
    """Create performance comparison chart"""
    if not results_history:
        return None
    
    df = pd.DataFrame(results_history)
    
    fig = go.Figure()
    
    fig.add_trace(go.Scatter(
        x=df.index,
        y=df['processing_time'],
        mode='lines+markers',
        name='Neural Model',
        line=dict(color='blue')
    ))
    
    fig.update_layout(
        title='Processing Time Comparison',
        xaxis_title='Request Number',
        yaxis_title='Processing Time (seconds)',
        height=400
    )
    
    return fig

# Header
st.markdown("""
<div class="main-header">
    <h1>🔍 Semantic Textual Similarity Application</h1>
    <h4>Universitas Mikroskil</h4>
    <p>Chrisandy | Fahim | Sandy</p>
    <p>2025 / 2026</p>
</div>
""", unsafe_allow_html=True)

# API Status Check
api_health = check_api_health()
api_connected = api_health["status"] in ["healthy", "unhealthy"]

# if api_connected:
#     if api_health["model_loaded"]:
#         st.markdown("""
#         <div class="api-status api-connected">
#             🟢 <strong>API Connected</strong> | Neural Model Loaded | Ready for Analysis!
#         </div>
#         """, unsafe_allow_html=True)
#     else:
#         st.markdown("""
#         <div class="api-status api-connected">
#             🟡 <strong>API Connected</strong> | Neural Model NOT Loaded
#         </div>
#         """, unsafe_allow_html=True)
# else:
#     st.markdown("""
#     <div class="api-status api-disconnected">
#         🔴 <strong>API Disconnected</strong> | Please start FastAPI backend first
#     </div>
#     """, unsafe_allow_html=True)

# Sidebar
st.sidebar.markdown("## 📋 Menu Navigasi")
analysis_type = st.sidebar.radio(
    "Pilih jenis analisis:",
    ["📄 Text Similarity", "📁 Document Similarity"],
    index=0
)

st.sidebar.markdown("## 🧠 Neural Model")
st.sidebar.markdown("**Embedding:** Paraphrase-MiniLM-L6-v2")
st.sidebar.markdown("**Single Encoder:** BiLSTM + Attention + MLP")

# Performance tracking
if 'performance_history' not in st.session_state:
    st.session_state.performance_history = []

# Main Application
if analysis_type == "📄 Text Similarity":
    st.markdown("## 🔍 Text Similarity Analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📝 Teks 1")
        text1 = st.text_area("Masukkan teks pertama (maksimal 50 kata):", height=200, key="text1", 
                            placeholder="Paste atau ketik teks pertama di sini...")
        
        if text1:
            word_count1 = len(text1.split())
            if word_count1 == 0:
                st.error("❌ Teks tidak boleh kosong")
                st.session_state.text1_valid = False
            elif word_count1 > 50:
                st.error(f"❌ Teks terlalu panjang: {word_count1} kata (maksimal 50 kata)")
                st.session_state.text1_valid = False
            else:
                is_valid1, msg1 = validate_document_length(text1)
                if is_valid1:
                    st.success(f"✅ {msg1} ({word_count1} kata)")
                    st.session_state.text1_valid = True
                else:
                    st.error(msg1)
                    st.session_state.text1_valid = False
        else:
            st.session_state.text1_valid = False
    
    with col2:
        st.markdown("### 📝 Teks 2")
        text2 = st.text_area("Masukkan teks kedua (maksimal 50 kata):", height=200, key="text2",
                            placeholder="Paste atau ketik teks kedua di sini...")
        
        if text2:
            word_count2 = len(text2.split())
            if word_count2 == 0:
                st.error("❌ Teks tidak boleh kosong")
                st.session_state.text2_valid = False
            elif word_count2 > 50:
                st.error(f"❌ Teks terlalu panjang: {word_count2} kata (maksimal 50 kata)")
                st.session_state.text2_valid = False
            else:
                is_valid2, msg2 = validate_document_length(text2)
                if is_valid2:
                    st.success(f"✅ {msg2} ({word_count2} kata)")
                    st.session_state.text2_valid = True
                else:
                    st.error(msg2)
                    st.session_state.text2_valid = False
        else:
            st.session_state.text2_valid = False
    
    # Analysis button
    # Cek apakah kedua teks valid
    both_texts_valid = st.session_state.get('text1_valid', False) and st.session_state.get('text2_valid', False)
    text_button_disabled = not (api_connected and api_health["model_loaded"] and both_texts_valid)
    
    # Tampilkan pesan helper jika ada yang belum valid
    if not both_texts_valid:
        if not st.session_state.get('text1_valid', False) and not st.session_state.get('text2_valid', False):
            st.info("ℹ️ Masukkan kedua teks yang valid untuk mulai analisis (maksimal 50 kata per teks)")
        elif not st.session_state.get('text1_valid', False):
            st.warning("⚠️ Teks 1 belum valid atau belum diisi")
        elif not st.session_state.get('text2_valid', False):
            st.warning("⚠️ Teks 2 belum valid atau belum diisi")
    
    col_center = st.columns([1, 2, 1])[1]
    with col_center:
        analyze_button = st.button(
            "🔍 Cek Plagiarisme / Kesamaan", 
            type="primary",
            disabled=text_button_disabled,
            use_container_width=True
        )
        
        if analyze_button:
            if not text1 or not text2:
                st.error("❌ Mohon masukkan kedua teks yang valid!")
            else:
                # Show loading
                with st.spinner("🔄 Menganalisis dengan Neural Model..."):
                    result = predict_similarity_api(text1, text2)
                
                err = result.get("error") or result.get("detail")
                if err:
                    st.error(f"❌ {err}")
                else:
                    similarity = result['similarity']
                    processing_time = result['processing_time']
                    method = result['method']
                    
                    # Add to performance history
                    st.session_state.performance_history.append({
                        'processing_time': processing_time,
                        'similarity': similarity
                    })
                    
                    # Results display with modern design
                    st.markdown("---")
                    st.markdown("""
                    <div style='text-align: center; padding: 1rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                                border-radius: 15px; margin-bottom: 2rem;'>
                        <h2 style='color: white; margin: 0;'>📊 Hasil Analisis Kesamaan</h2>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Main Score Card - Large and Centered
                    col_space1, col_main, col_space2 = st.columns([1, 2, 1])
                    with col_main:
                        if similarity >= 0.8:
                            bg_color = "linear-gradient(135deg, #11998e 0%, #38ef7d 100%)"
                            status_icon = "🟢"
                            status_text = "SANGAT MIRIP"
                            status_desc = "Terdeteksi kemiripan sangat tinggi"
                        elif similarity >= 0.6:
                            bg_color = "linear-gradient(135deg, #f093fb 0%, #f5576c 100%)"
                            status_icon = "🟡"
                            status_text = "CUKUP MIRIP"
                            status_desc = "Terdeteksi kemiripan sedang"
                        else:
                            bg_color = "linear-gradient(135deg, #4facfe 0%, #00f2fe 100%)"
                            status_icon = "🔵"
                            status_text = "BERBEDA"
                            status_desc = "Tidak terdeteksi kemiripan signifikan"
                        
                        st.markdown(f"""
                        <div style='background: {bg_color}; padding: 2.5rem; border-radius: 20px; 
                                    text-align: center; box-shadow: 0 10px 30px rgba(0,0,0,0.3);'>
                            <div style='font-size: 4rem; margin-bottom: 0.5rem;'>{status_icon}</div>
                            <div style='color: white; font-size: 3.5rem; font-weight: 800; margin-bottom: 0.5rem;'>
                                {similarity*100:.1f}%
                            </div>
                            <div style='color: rgba(255,255,255,0.9); font-size: 1.2rem; font-weight: 600; 
                                        letter-spacing: 2px; margin-bottom: 0.5rem;'>
                                {status_text}
                            </div>
                            <div style='color: rgba(255,255,255,0.8); font-size: 0.95rem;'>
                                {status_desc}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    # Performance metrics with modern cards
                    perf_cols = st.columns(3)
                    
                    with perf_cols[0]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #667eea;'>
                            <div style='color: #667eea; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                📏 SIMILARITY SCORE
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #2d3748;'>
                                {similarity:.4f}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with perf_cols[1]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #764ba2;'>
                            <div style='color: #764ba2; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                ⚡ PROCESSING TIME
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #2d3748;'>
                                {processing_time:.3f}s
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with perf_cols[2]:
                        model_status = "Trained" if result.get('weights_loaded') else "Untrained"
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #667eea;'>
                            <div style='color: #667eea; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                🤖 MODEL STATUS
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #2d3748;'>
                                {model_status}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    # Gauge visualization
                    gauge_fig = create_similarity_gauge(similarity, method, processing_time)
                    st.plotly_chart(gauge_fig, use_container_width=True)
                    
                    # Text statistics
                    st.markdown("## 📈 Text Statistics")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("#### 📄 Teks 1")
                        stats1 = get_text_statistics(text1)
                        for key, value in stats1.items():
                            st.metric(key.replace('_', ' ').title(), value)
                    
                    with col2:
                        st.markdown("#### 📄 Teks 2") 
                        stats2 = get_text_statistics(text2)
                        for key, value in stats2.items():
                            st.metric(key.replace('_', ' ').title(), value)

elif analysis_type == "📁 Document Similarity":
    st.markdown("## 📁 Document Similarity Analysis")
    
    # Info message based on available support
    supported_formats = []
    if PDF_SUPPORT:
        supported_formats.append("PDF")
    if DOCX_SUPPORT:
        supported_formats.append("DOCX")
    supported_formats.append("TXT")
    
    if not PDF_SUPPORT and not DOCX_SUPPORT:
        st.warning("📦 **PDF & DOCX support tidak tersedia.** Install PyPDF2, pdfplumber, dan python-docx untuk support lengkap.")
        st.info("🔧 Upload tepat 2 dokumen TXT untuk dibandingkan (maksimal 5 halaman per dokumen)")
    else:
        st.info(f"🔧 Upload tepat 2 dokumen ({', '.join(supported_formats)}) untuk dibandingkan (maksimal 5 halaman per dokumen)")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📄 Dokumen 1")
        file_types = ["txt"]
        if PDF_SUPPORT:
            file_types.append("pdf")
        if DOCX_SUPPORT:
            file_types.extend(["docx", "doc"])
        
        uploaded_file1 = st.file_uploader(
            f"Upload dokumen pertama ({', '.join([t.upper() for t in file_types])})", 
            type=file_types, 
            key="file1"
        )
        doc1_content = ""
        
        if uploaded_file1:
            # Handle DOCX files
            if uploaded_file1.name.lower().endswith(('.docx', '.doc')):
                st.session_state.doc1_valid = False
                if DOCX_SUPPORT:
                    is_valid_docx, docx_msg = validate_docx_file(uploaded_file1)
                    if is_valid_docx:
                        with st.spinner("📖 Mengekstrak teks dari DOCX..."):
                            text, error = extract_text_from_docx(uploaded_file1)
                        if error:
                            st.error(error)
                            st.session_state.doc1_valid = False
                        else:
                            doc1_content = text
                            is_valid1, msg1 = validate_document_length(doc1_content)
                            
                            if is_valid1:
                                st.success("✅ DOCX berhasil diproses!")
                                st.success(msg1)
                                with st.expander("👁️ Preview Dokumen 1 (dari DOCX)"):
                                    st.text_area("Content:", create_docx_preview(doc1_content), 
                                               height=150, disabled=True, key="preview1_docx")
                                st.session_state.doc1_valid = True
                            else:
                                st.error(msg1)
                                doc1_content = ""
                                st.session_state.doc1_valid = False
                    else:
                        st.error(docx_msg)
                        st.session_state.doc1_valid = False
                else:
                    st.error("❌ DOCX support tidak tersedia")
                    st.session_state.doc1_valid = False
            
            # Handle PDF files
            elif uploaded_file1.type == "application/pdf":
                st.session_state.doc1_valid = False
                if PDF_SUPPORT:
                    is_valid_pdf, pdf_msg = validate_pdf_file(uploaded_file1)
                    if is_valid_pdf:
                        with st.spinner("📖 Mengekstrak teks dari PDF..."):
                            text, error = extract_text_from_pdf(uploaded_file1)
                        if error:
                            st.error(error)
                            st.session_state.doc1_valid = False
                        else:
                            doc1_content = text
                            is_valid1, msg1 = validate_document_length(doc1_content)
                            
                            if is_valid1:
                                st.success("✅ PDF berhasil diproses!")
                                st.success(msg1)
                                with st.expander("👁️ Preview Dokumen 1 (dari PDF)"):
                                    st.text_area("Content:", create_pdf_preview(doc1_content), 
                                               height=150, disabled=True, key="preview1")
                                st.session_state.doc1_valid = True
                            else:
                                st.error(msg1)
                                doc1_content = ""
                                st.session_state.doc1_valid = False
                    else:
                        st.error(pdf_msg)
                        st.session_state.doc1_valid = False
                else:
                    st.error("❌ PDF support tidak tersedia")
                    st.session_state.doc1_valid = False
            
            # Handle TXT files
            else:
                st.session_state.doc1_valid = False
                # TXT file
                doc1_content = uploaded_file1.read().decode("utf-8")
                is_valid1, msg1 = validate_document_length(doc1_content)
                
                if is_valid1:
                    st.success("✅ File TXT berhasil dimuat!")
                    st.success(msg1)
                    with st.expander("👁️ Preview Dokumen 1"):
                        st.text_area("Content:", doc1_content[:500] + "..." if len(doc1_content) > 500 else doc1_content, 
                                   height=150, disabled=True, key="preview1_txt")
                    st.session_state.doc1_valid = True
                else:
                    st.error(msg1)
                    doc1_content = ""
                    st.session_state.doc1_valid = False
        else:
            st.session_state.doc1_valid = False
    
    with col2:
        st.markdown("### 📄 Dokumen 2")
        uploaded_file2 = st.file_uploader(
            f"Upload dokumen kedua ({', '.join([t.upper() for t in file_types])})", 
            type=file_types, 
            key="file2"
        )
        doc2_content = ""
        
        if uploaded_file2:
            # Handle DOCX files
            if uploaded_file2.name.lower().endswith(('.docx', '.doc')):
                st.session_state.doc2_valid = False
                if DOCX_SUPPORT:
                    is_valid_docx, docx_msg = validate_docx_file(uploaded_file2)
                    if is_valid_docx:
                        with st.spinner("📖 Mengekstrak teks dari DOCX..."):
                            text, error = extract_text_from_docx(uploaded_file2)
                        if error:
                            st.error(error)
                            st.session_state.doc2_valid = False
                        else:
                            doc2_content = text
                            is_valid2, msg2 = validate_document_length(doc2_content)
                            
                            if is_valid2:
                                st.success("✅ DOCX berhasil diproses!")
                                st.success(msg2)
                                with st.expander("👁️ Preview Dokumen 2 (dari DOCX)"):
                                    st.text_area("Content:", create_docx_preview(doc2_content), 
                                               height=150, disabled=True, key="preview2_docx")
                                st.session_state.doc2_valid = True
                            else:
                                st.error(msg2)
                                doc2_content = ""
                                st.session_state.doc2_valid = False
                    else:
                        st.error(docx_msg)
                        st.session_state.doc2_valid = False
                else:
                    st.error("❌ DOCX support tidak tersedia")
                    st.session_state.doc2_valid = False
            
            # Handle PDF files
            elif uploaded_file2.type == "application/pdf":
                st.session_state.doc2_valid = False
                if PDF_SUPPORT:
                    is_valid_pdf, pdf_msg = validate_pdf_file(uploaded_file2)
                    if is_valid_pdf:
                        with st.spinner("📖 Mengekstrak teks dari PDF..."):
                            text, error = extract_text_from_pdf(uploaded_file2)
                        if error:
                            st.error(error)
                            st.session_state.doc2_valid = False
                        else:
                            doc2_content = text
                            is_valid2, msg2 = validate_document_length(doc2_content)
                            
                            if is_valid2:
                                st.success("✅ PDF berhasil diproses!")
                                st.success(msg2)
                                with st.expander("👁️ Preview Dokumen 2 (dari PDF)"):
                                    st.text_area("Content:", create_pdf_preview(doc2_content), 
                                               height=150, disabled=True, key="preview2")
                                st.session_state.doc2_valid = True
                            else:
                                st.error(msg2)
                                doc2_content = ""
                                st.session_state.doc2_valid = False
                    else:
                        st.error(pdf_msg)
                        st.session_state.doc2_valid = False
                else:
                    st.error("❌ PDF support tidak tersedia")
                    st.session_state.doc2_valid = False
            
            # Handle TXT files
            else:
                st.session_state.doc2_valid = False
                # TXT file
                doc2_content = uploaded_file2.read().decode("utf-8")
                is_valid2, msg2 = validate_document_length(doc2_content)
                
                if is_valid2:
                    st.success("✅ File TXT berhasil dimuat!")
                    st.success(msg2)
                    with st.expander("👁️ Preview Dokumen 2"):
                        st.text_area("Content:", doc2_content[:500] + "..." if len(doc2_content) > 500 else doc2_content, 
                                   height=150, disabled=True, key="preview2_txt")
                    st.session_state.doc2_valid = True
                else:
                    st.error(msg2)
                    doc2_content = ""
                    st.session_state.doc2_valid = False
        else:
            st.session_state.doc2_valid = False
    
    # Analysis button
    # Cek apakah kedua dokumen valid
    both_docs_valid = st.session_state.get('doc1_valid', False) and st.session_state.get('doc2_valid', False)
    button_disabled = not (api_connected and api_health["model_loaded"] and both_docs_valid)
    
    # Tampilkan pesan helper jika ada yang belum valid
    if not both_docs_valid:
        if not st.session_state.get('doc1_valid', False) and not st.session_state.get('doc2_valid', False):
            st.info("ℹ️ Upload kedua dokumen yang valid untuk mulai analisis")
        elif not st.session_state.get('doc1_valid', False):
            st.warning("⚠️ Dokumen 1 belum valid atau belum di-upload")
        elif not st.session_state.get('doc2_valid', False):
            st.warning("⚠️ Dokumen 2 belum valid atau belum di-upload")
    
    if st.button("🔍 Cek Plagiarisme / Kesamaan", type="primary", disabled=button_disabled, use_container_width=True):
        if not doc1_content or not doc2_content:
            st.error("❌ Mohon upload kedua dokumen yang valid!")
        else:
            # Pastikan model neural tersedia, karena /predict-document butuh model neural
            if not api_health.get("model_loaded"):
                st.error("🟡 Neural model belum loaded di backend. Buka /health di FastAPI dan pastikan artifacts lengkap. (/predict-document butuh model neural)")
            else:
                with st.spinner("📊 Menganalisis dokumen (sliding window + BERTScore)..."):
                    # Kamu bisa expose parameter di sidebar kalau mau
                    result = predict_document_api(
                        doc1_content, doc2_content,
                        per_side_len=28,  # samakan dengan backend default
                        stride=28,
                        topk_evidence=5,
                        use_symmetric=True
                    )
                if "error" in result:
                    st.error(f"❌ {result['error']}")
                else:
                    doc_score = result["doc_score"]
                    processing_time = result["processing_time"]
                    detail = result["detail"]
                    top_evidence = result["top_evidence"]
                    shape = result["shape"]

                    # Modern Results Display for Document
                    st.markdown("---")
                    st.markdown("""
                    <div style='text-align: center; padding: 1rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                                border-radius: 15px; margin-bottom: 2rem;'>
                        <h2 style='color: white; margin: 0;'>🎯 Hasil Analisis Dokumen</h2>
                        <p style='color: rgba(255,255,255,0.9); margin: 0.5rem 0 0 0; font-size: 0.9rem;'>BERTScore-like F1 Analysis</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Main Score Card
                    col_space1, col_main, col_space2 = st.columns([1, 2, 1])
                    with col_main:
                        if doc_score >= 0.8:
                            bg_color = "linear-gradient(135deg, #f093fb 0%, #f5576c 100%)"
                            status_icon = "🚨"
                            status_text = "PLAGIARISME TERDETEKSI"
                            status_desc = "Kemiripan sangat tinggi - potensi plagiarisme"
                        elif doc_score >= 0.6:
                            bg_color = "linear-gradient(135deg, #ffeaa7 0%, #fdcb6e 100%)"
                            status_icon = "⚠️"
                            status_text = "PERLU DITINJAU"
                            status_desc = "Kemiripan cukup tinggi - perlu verifikasi"
                        else:
                            bg_color = "linear-gradient(135deg, #11998e 0%, #38ef7d 100%)"
                            status_icon = "✅"
                            status_text = "DOKUMEN UNIK"
                            status_desc = "Tidak terdeteksi plagiarisme signifikan"
                        
                        st.markdown(f"""
                        <div style='background: {bg_color}; padding: 2.5rem; border-radius: 20px; 
                                    text-align: center; box-shadow: 0 10px 30px rgba(0,0,0,0.3);'>
                            <div style='font-size: 4rem; margin-bottom: 0.5rem;'>{status_icon}</div>
                            <div style='color: white; font-size: 3.5rem; font-weight: 800; margin-bottom: 0.5rem;'>
                                {doc_score*100:.1f}%
                            </div>
                            <div style='color: rgba(255,255,255,0.9); font-size: 1.2rem; font-weight: 600; 
                                        letter-spacing: 2px; margin-bottom: 0.5rem;'>
                                {status_text}
                            </div>
                            <div style='color: rgba(255,255,255,0.8); font-size: 0.95rem;'>
                                {status_desc}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    # Performance metrics with modern cards
                    perf_cols = st.columns(3)
                    
                    with perf_cols[0]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #667eea;'>
                            <div style='color: #667eea; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                📊 F1 SCORE
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #2d3748;'>
                                {doc_score:.4f}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with perf_cols[1]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #764ba2;'>
                            <div style='color: #764ba2; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                📏 WINDOW PAIRS
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #2d3748;'>
                                {shape['m']} × {shape['n']}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with perf_cols[2]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #667eea;'>
                            <div style='color: #667eea; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                ⚡ PROCESSING TIME
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #2d3748;'>
                                {processing_time:.3f}s
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("<br>", unsafe_allow_html=True)

                    # Gauge pakai doc_score
                    gauge_fig = create_similarity_gauge(doc_score, "Sliding Window + BERTScore", processing_time)
                    st.plotly_chart(gauge_fig, use_container_width=True)

                    # Detail P/R/F1 with modern cards
                    st.markdown("""
                    <div style='background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); 
                                padding: 0.75rem; border-radius: 10px; margin: 1.5rem 0;'>
                        <h3 style='color: white; margin: 0; font-size: 1.1rem;'>📌 Precision / Recall / F1 Metrics</h3>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if detail.get("symmetric", False):
                        a2b, b2a = detail["AtoB"], detail["BtoA"]
                        metric_cols = st.columns(2)
                        
                        with metric_cols[0]:
                            st.markdown(f"""
                            <div style='background: rgba(102, 126, 234, 0.05); padding: 1.5rem; border-radius: 12px; border: 2px solid #667eea;'>
                                <div style='color: #667eea; font-weight: 600; margin-bottom: 1rem; font-size: 1.1rem;'>📄 Dokumen A → B</div>
                                <div style='display: flex; justify-content: space-between; margin-bottom: 0.5rem;'>
                                    <span style='color: #4a5568;'>Precision:</span>
                                    <span style='font-weight: 700; color: #2d3748;'>{a2b['P']:.3f}</span>
                                </div>
                                <div style='display: flex; justify-content: space-between; margin-bottom: 0.5rem;'>
                                    <span style='color: #4a5568;'>Recall:</span>
                                    <span style='font-weight: 700; color: #2d3748;'>{a2b['R']:.3f}</span>
                                </div>
                                <div style='display: flex; justify-content: space-between; padding-top: 0.5rem; border-top: 2px solid #667eea;'>
                                    <span style='color: #667eea; font-weight: 600;'>F1 Score:</span>
                                    <span style='font-weight: 800; color: #667eea; font-size: 1.3rem;'>{a2b['F1']:.3f}</span>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                        
                        with metric_cols[1]:
                            st.markdown(f"""
                            <div style='background: rgba(118, 75, 162, 0.05); padding: 1.5rem; border-radius: 12px; border: 2px solid #764ba2;'>
                                <div style='color: #764ba2; font-weight: 600; margin-bottom: 1rem; font-size: 1.1rem;'>📄 Dokumen B → A</div>
                                <div style='display: flex; justify-content: space-between; margin-bottom: 0.5rem;'>
                                    <span style='color: #4a5568;'>Precision:</span>
                                    <span style='font-weight: 700; color: #2d3748;'>{b2a['P']:.3f}</span>
                                </div>
                                <div style='display: flex; justify-content: space-between; margin-bottom: 0.5rem;'>
                                    <span style='color: #4a5568;'>Recall:</span>
                                    <span style='font-weight: 700; color: #2d3748;'>{b2a['R']:.3f}</span>
                                </div>
                                <div style='display: flex; justify-content: space-between; padding-top: 0.5rem; border-top: 2px solid #764ba2;'>
                                    <span style='color: #764ba2; font-weight: 600;'>F1 Score:</span>
                                    <span style='font-weight: 800; color: #764ba2; font-size: 1.3rem;'>{b2a['F1']:.3f}</span>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                    else:
                        a2b = detail["AtoB"]
                        st.markdown(f"""
                        <div style='background: rgba(102, 126, 234, 0.05); padding: 1.5rem; border-radius: 12px; border: 2px solid #667eea;'>
                            <div style='color: #667eea; font-weight: 600; margin-bottom: 1rem; font-size: 1.1rem;'>📄 Dokumen A → B</div>
                            <div style='display: flex; justify-content: space-between; margin-bottom: 0.5rem;'>
                                <span style='color: #4a5568;'>Precision:</span>
                                <span style='font-weight: 700; color: #2d3748;'>{a2b['P']:.3f}</span>
                            </div>
                            <div style='display: flex; justify-content: space-between; margin-bottom: 0.5rem;'>
                                <span style='color: #4a5568;'>Recall:</span>
                                <span style='font-weight: 700; color: #2d3748;'>{a2b['R']:.3f}</span>
                            </div>
                            <div style='display: flex; justify-content: space-between; padding-top: 0.5rem; border-top: 2px solid #667eea;'>
                                <span style='color: #667eea; font-weight: 600;'>F1 Score:</span>
                                <span style='font-weight: 800; color: #667eea; font-size: 1.3rem;'>{a2b['F1']:.3f}</span>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                    # Top evidence windows with modern design
                    st.markdown("""
                    <div style='background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); 
                                padding: 0.75rem; border-radius: 10px; margin: 1.5rem 0;'>
                        <h3 style='color: white; margin: 0; font-size: 1.1rem;'>🔎 Top Evidence Windows</h3>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if top_evidence:
                        for k, ev in enumerate(top_evidence, 1):
                            score_color = "#f5576c" if ev['score'] >= 0.8 else "#fdcb6e" if ev['score'] >= 0.6 else "#38ef7d"
                            st.markdown(f"""
                            <div style='background: rgba(102, 126, 234, 0.03); padding: 1.25rem; border-radius: 12px; 
                                        margin-bottom: 1rem; border-left: 5px solid {score_color};'>
                                <div style='display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.75rem;'>
                                    <span style='font-weight: 700; color: #2d3748; font-size: 1.1rem;'>Evidence #{k}</span>
                                    <span style='background: {score_color}; color: white; padding: 0.25rem 0.75rem; 
                                                border-radius: 20px; font-weight: 700; font-size: 0.9rem;'>
                                        {ev['score']:.3f}
                                    </span>
                                </div>
                                <div style='background: white; padding: 0.75rem; border-radius: 8px; margin-bottom: 0.5rem; 
                                            border: 1px solid #e2e8f0;'>
                                    <div style='color: #667eea; font-weight: 600; font-size: 0.85rem; margin-bottom: 0.25rem;'>
                                        📄 Window A:
                                    </div>
                                    <div style='color: #2d3748; line-height: 1.6;'>{ev["windowA"]}</div>
                                </div>
                                <div style='background: white; padding: 0.75rem; border-radius: 8px; border: 1px solid #e2e8f0;'>
                                    <div style='color: #764ba2; font-weight: 600; font-size: 0.85rem; margin-bottom: 0.25rem;'>
                                        📄 Window B:
                                    </div>
                                    <div style='color: #2d3748; line-height: 1.6;'>{ev["windowB"]}</div>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                    else:
                        st.info("ℹ️ Tidak ada evidence (kemungkinan dokumen sangat pendek).")

# Footer & Additional Info
st.markdown("---")

# Performance Analytics
if st.session_state.performance_history:
    with st.expander("📊 Performance Analytics"):
        perf_chart = create_performance_chart(st.session_state.performance_history)
        if perf_chart:
            st.plotly_chart(perf_chart, width="stretch")
        
        # Performance summary
        df_perf = pd.DataFrame(st.session_state.performance_history)
        st.markdown("### Performance Summary:")
        
        if len(df_perf) > 0:
            avg_time = df_perf['processing_time'].mean()
            avg_similarity = df_perf['similarity'].mean()
            
            summary_cols = st.columns(2)
            with summary_cols[0]:
                st.metric("Average Time", f"{avg_time:.3f}s")
            with summary_cols[1]:
                st.metric("Average Similarity", f"{avg_similarity:.3f}")
